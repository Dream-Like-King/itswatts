from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "Its_Watts_Release_1_Requirements.docx"

# compact_reference_guide tokens, with a named It’s Watts green accent override.
INK = RGBColor(23, 35, 27)
GREEN = RGBColor(22, 132, 58)
MUTED = RGBColor(89, 102, 93)
PALE_GREEN = "E8F3EB"


def font(run, size, color=INK, bold=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def shade(paragraph, fill=PALE_GREEN):
    props = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    props.append(node)


def bottom_rule(paragraph, color="D5DDD7"):
    props = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "7")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    props.append(borders)


def page_number(paragraph):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph.add_run()._r.addnext(field)


def body(doc, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text), 11)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    font(p.add_run(text), 16 if level == 1 else 12.5, GREEN if level == 1 else INK, True)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text), 11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text), 11)
    return p


def label_bullet(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(f"{label}. "), 11, INK, True)
    font(p.add_run(text), 11)
    return p


def callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.14)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(10)
    shade(p)
    font(p.add_run(f"{label}: "), 10.5, GREEN, True)
    font(p.add_run(text), 10.5)
    return p


def phase(doc, number, name, goal, required, acceptance):
    heading(doc, f"Phase {number}: {name}", 2)
    body(doc, goal, after=4)
    label_bullet(doc, "Required", required)
    label_bullet(doc, "Acceptance", acceptance)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.49)
section.footer_distance = Inches(0.49)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)

header = section.header.paragraphs[0]
header.paragraph_format.space_after = Pt(4)
font(header.add_run("IT’S WATTS  |  RELEASE 1 REQUIREMENTS"), 8.5, MUTED, True)
bottom_rule(header)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(footer.add_run("ITSWATTS.COM  |  PAGE "), 8.5, MUTED)
page_number(footer)

# Editorial-cover header pattern.
doc.add_paragraph().paragraph_format.space_after = Pt(36)
kicker = doc.add_paragraph()
kicker.paragraph_format.space_after = Pt(8)
font(kicker.add_run("PRODUCT REQUIREMENTS DOCUMENT"), 9, GREEN, True)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(8)
font(title.add_run("It’s Watts\nRelease 1"), 29, INK, True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(15)
font(subtitle.add_run("Requirements for the original Phase 1–10 site launch"), 14, GREEN, True)

meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(30)
font(meta.add_run("Version 1.0  •  July 2026  •  Owner: Derrick Watson"), 10.5, MUTED)
bottom_rule(meta, "16843A")

callout(doc, "Release decision", "Release 1 establishes It’s Watts as a concise QA education hub—not a traditional portfolio. The launch combines practical QA resources, small interactive tools, thoughtful AI support, and a polished, accessible experience.")

heading(doc, "1. Product purpose and release scope")
body(doc, "Release 1 creates the foundation for itswatts.com. It is a public, responsive website focused on QA education, automation, AI-assisted quality work, and practical resources. The release must feel useful to learners and professionals while keeping the homepage visually focused and easy to navigate on mobile.")
label_bullet(doc, "Primary audience", "Early-career QA professionals, testers moving into automation, developers strengthening quality habits, and teams seeking concise QA resources.")
label_bullet(doc, "Product promise", "Useful QA education for the AI era: clearer thinking, stronger testing habits, and software people can trust.")
label_bullet(doc, "Release boundary", "Release 1 provides static educational content, browser-based tools, downloadable resources, and an optional AI chat integration. User accounts, saved progress, payment, community moderation, and a CMS are outside scope.")

heading(doc, "2. Release 1 functional requirements")
phase(doc, "1", "Foundation and brand entry point", "Establish a production-ready React, Vite, and TypeScript website with a cohesive It’s Watts identity.", "Build a responsive application structure with reusable components, a glass-style navigation bar, an animated hero, clear calls to action, and the written brand name It’s Watts. The logo may display it’s wattϟ, with the bolt acting as the stylized final ‘s.’", "Visitors can load the home page on desktop or mobile, understand the QA education focus, use the navigation, and reach the primary content without broken layout or branding inconsistencies.")
phase(doc, "2", "QA education positioning", "Move the primary experience beyond a personal portfolio toward practical QA education.", "Use concise content that explains the site’s focus on quality engineering, automation, AI, accessibility, and practical learning. Position the portfolio and professional background as supporting context rather than the main destination.", "A first-time visitor can identify that It’s Watts is a QA education and toolkit site within the opening screen and can choose a learning or tool destination.")
phase(doc, "3", "Learning paths", "Give visitors focused ways to explore QA topics without reading a long, unstructured feed.", "Provide compact paths for Automation Foundations, AI for QA, and Quality Essentials. Each path must explain its purpose, list focused topic tags, and open deeper content in an accessible dialog or focused view.", "Learning-path cards are readable on mobile, open and close reliably, support keyboard use, and avoid putting all lesson content directly on the homepage.")
phase(doc, "4", "QA toolkits", "Provide practical tools that help visitors make better quality decisions.", "Include an automation decision guide, a test-case starter that changes output based on the visitor’s description, and a QA prompt library with copyable examples. Tools must work without requiring a login.", "A visitor can interact with every tool using mouse, touch, or keyboard; generated test-case suggestions reflect the entered scenario; and tool outputs remain readable in both themes.")
phase(doc, "5", "Weekly learning and resources", "Create a sustainable content surface for notes, templates, and reusable QA material.", "Include Weekly Notes with readable article dialogs, downloadable QA templates, and a compact Resource Library that reveals deeper material only when requested. Resource content must not turn the homepage into an endless mobile scroll.", "Visitors can open a note, browse resource details, download available templates, and return to the main site without losing context.")
phase(doc, "6", "Ask Watt chatbot", "Offer an AI assistant experience focused on QA education and the site’s practical resources.", "Provide the Ask Watt user interface and a server-side API route that accepts questions, keeps OPENAI_API_KEY outside browser code, and handles unavailable billing or service errors with a friendly message.", "The chat interface is usable without exposing secrets. When a valid key and API credits are configured, questions reach the server endpoint; when unavailable, the site explains the temporary limitation without breaking the page.")
phase(doc, "7", "Visual system and brand refinement", "Create a premium technical toolkit visual language that remains readable and purposeful.", "Use the established dark charcoal foundation, subtly elevated panels, thin technical borders, electric green accents, and clear typography. Include a matching light theme rather than a separate visual identity.", "Panels remain distinct from the page background, text has sufficient contrast, written references say It’s Watts, and the logo treatment is consistent across navigation and footer.")
phase(doc, "8", "Accessibility and responsive experience", "Make the site comfortable to use across screen sizes, themes, and input methods.", "Support system light/dark preference, visible focus indicators, keyboard navigation, a mobile menu with adequate contrast, reduced-motion behavior, semantic labels, and no required horizontal scrolling.", "Key content, dialogs, tools, and navigation are usable on narrow mobile screens and in both themes; text remains legible; and interaction is possible without a mouse.")
phase(doc, "9", "Discoverability and sharing", "Prepare the public site for search engines and professional sharing.", "Add page title and description metadata, canonical URL, Open Graph and social-image metadata, a branded share image, robots.txt, sitemap.xml, and correct public URLs.", "A deployed page presents the expected title, description, and social preview references. Search-engine support files are publicly available and point to the production domain.")
phase(doc, "10", "Launch polish and content management", "Complete Release 1 with maintainable content patterns and production verification.", "Keep weekly notes and resource data easy to update in the codebase, preserve a compact home-page layout, verify build and lint checks, and configure deployment through Vercel with environment-variable support for Ask Watt.", "Production builds and lint checks pass. Content updates can be made through the site’s data and public-download files, and deployment can occur without placing private API keys in the repository.")

heading(doc, "3. Non-functional requirements")
for item in [
    "Mobile-first experience: no horizontal scrolling; full resource details should open in dialogs or dedicated views rather than extending the home page indefinitely.",
    "Accessibility: maintain readable contrast, semantic controls, keyboard support, focus visibility, Escape-to-close behavior for dialogs, and reduced-motion support.",
    "Theme consistency: all pages, dialogs, navigation, tools, and content must work in both system light and dark modes.",
    "Performance: favor lightweight CSS and browser-native behavior; avoid large unoptimized media or unnecessary client-side dependencies.",
    "Privacy and security: never expose OPENAI_API_KEY in the client bundle, source repository, or browser network payload. Store it only in deployment environment settings.",
    "Maintainability: use reusable React components and centralized content data where possible so notes, paths, and resources can be updated without redesigning the site.",
]:
    bullet(doc, item)

heading(doc, "4. Vercel deployment requirements")
body(doc, "Vercel is required for the public deployment workflow because Release 1 combines a static Vite website with a lightweight server-side Ask Watt endpoint.")
for item in [
    "Deploy the website from the connected source repository with production and preview deployments.",
    "Serve public static content, images, and downloadable templates efficiently.",
    "Run the Ask Watt API route server-side so the OpenAI key stays private.",
    "Store OPENAI_API_KEY as a Vercel environment variable; OpenAI API billing remains separate from Vercel hosting.",
    "Support the custom itswatts.com domain and HTTPS delivery.",
]:
    bullet(doc, item)

heading(doc, "5. Release acceptance checklist")
for item in [
    "The homepage communicates QA education, automation, AI, and tools—not only a personal portfolio.",
    "Navigation, learning paths, toolkits, Weekly Notes, resources, and contact surfaces work on desktop and mobile.",
    "The automation guide, test-case starter, and prompt library are functional and understandable.",
    "Ask Watt presents a secure, helpful experience and degrades gracefully if API billing is unavailable.",
    "Light and dark modes are readable, including dialogs, note content, learning-path content, and mobile navigation.",
    "All written brand references use It’s Watts; the logo uses it’s wattϟ intentionally.",
    "Social metadata, share image, robots.txt, sitemap.xml, and canonical domain settings are present.",
    "The production build and lint checks complete successfully before deployment.",
]:
    numbered(doc, item)

heading(doc, "6. Release 1 handoff and future boundary")
callout(doc, "Handoff", "Release 1 is the stable educational foundation. Future releases may add dedicated learning pages, a Knowledge Base, career-path guidance, a QA orientation guide, and the practical learning tools described in the future-updates requirements document. They should expand through focused destinations, not an endlessly longer homepage.")
body(doc, "Release 1 maintenance responsibilities", after=4)
for item in [
    "Publish new Weekly Notes by updating the central content data and reviewing both light and dark presentation before deployment.",
    "Add new downloads through the public resources area, using clear names and a short explanation of when each template is useful.",
    "Review Ask Watt after API billing, model, or prompt updates; keep its answers focused on QA education and do not expose private configuration.",
    "Use Vercel preview deployments to review meaningful changes before promoting them to the public domain.",
]:
    bullet(doc, item)
callout(doc, "Not included", "Release 1 does not require user accounts, saved learning progress, a newsletter platform, public submissions, paid content, a community forum, or a CMS. Those decisions belong to future planning after the initial education experience is established.")

doc.core_properties.title = "It’s Watts — Release 1 Requirements"
doc.core_properties.subject = "Original Phase 1–10 requirements"
doc.core_properties.author = "Derrick Watson"
doc.core_properties.keywords = "It’s Watts, QA education, Release 1, requirements"
doc.save(OUT)
print(OUT)
