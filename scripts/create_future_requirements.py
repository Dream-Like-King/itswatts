from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "Its_Watts_Future_Updates_Requirements.docx"

INK = RGBColor(23, 35, 27)
GREEN = RGBColor(22, 132, 58)
MUTED = RGBColor(89, 102, 93)
PALE_GREEN = "E8F3EB"


def set_font(run, size=None, color=None, bold=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def shade(paragraph, fill):
    props = paragraph._p.get_or_add_pPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    props.append(element)


def border_bottom(paragraph, color="D5DDD7"):
    props = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "7")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    props.append(borders)


def page_number(paragraph):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph.add_run()._r.addnext(field)


def para(doc, text, after=6, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), 11, color)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(7)
    set_font(p.add_run(text), 16, GREEN, True)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(text), 12.5, INK, True)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), 11, INK)
    return p


def label_bullet(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(f"{label}. "), 11, INK, True)
    set_font(p.add_run(text), 11, INK)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), 11, INK)
    return p


def callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.14)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(10)
    shade(p, PALE_GREEN)
    set_font(p.add_run(f"{label}: "), 10.5, GREEN, True)
    set_font(p.add_run(text), 10.5, INK)
    return p


def requirement(doc, number, name, purpose, requirements, acceptance):
    h2(doc, f"FR-{number}: {name}")
    para(doc, purpose, after=4)
    label_bullet(doc, "Required", requirements)
    label_bullet(doc, "Acceptance", acceptance)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.49)
section.footer_distance = Inches(0.49)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)

header = section.header.paragraphs[0]
header.paragraph_format.space_after = Pt(4)
set_font(header.add_run("IT’S WATTS  |  FUTURE UPDATES REQUIREMENTS"), 8.5, MUTED, True)
border_bottom(header)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.add_run("ITSWATTS.COM  |  PAGE "), 8.5, MUTED)
page_number(footer)

# Editorial-cover opening using the compact reference guide preset with brand-green override.
doc.add_paragraph().paragraph_format.space_after = Pt(36)
kicker = doc.add_paragraph()
kicker.paragraph_format.space_after = Pt(8)
set_font(kicker.add_run("PRODUCT REQUIREMENTS DOCUMENT"), 9, GREEN, True)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(8)
set_font(title.add_run("It’s Watts\nFuture Learning Platform Updates"), 28, INK, True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(15)
set_font(subtitle.add_run("Requirements for a concise, mobile-first QA education experience"), 14, GREEN, True)

metadata = doc.add_paragraph()
metadata.paragraph_format.space_after = Pt(30)
set_font(metadata.add_run("Version 1.0  •  July 2026  •  Owner: Derrick Watson"), 10.5, MUTED)
border_bottom(metadata, "16843A")

callout(doc, "Product decision", "Future education content must expand through focused pages, dialogs, and short previews—not by turning the homepage into a long scrolling feed. On mobile, each home section should help visitors choose a path quickly.")

h1(doc, "1. Purpose and product goal")
para(doc, "This document defines the next set of educational features for It’s Watts. The product goal is to help visitors understand QA, find the right learning level, build a practical career path, and access useful resources while preserving the site’s focused, premium, mobile-friendly experience.")
label_bullet(doc, "Primary audience", "People new to QA, early-career testers, testers moving into automation, and working professionals strengthening quality strategy.")
label_bullet(doc, "Primary value", "Clear next steps, practical explanations, and learning resources that are easier to act on than a long, unstructured content feed.")
label_bullet(doc, "Design principle", "Show a concise preview on the homepage; place deeper material on dedicated pages or inside intentional dialogs.")

h1(doc, "2. Scope and information architecture")
para(doc, "The homepage remains a short entry point. Future content should be organized into dedicated destinations, each reachable from clear navigation and compact home-page cards.")
for item in [
    "Learn: What Is QA?, QA in the SDLC, and guided beginner-to-advanced learning paths.",
    "Knowledge Base: searchable QA terms, definitions, and links to related lessons or tools.",
    "Career Paths: role progression, certification guidance, and the QA orientation guide.",
    "Resources: templates, downloadable guides, and the existing resource library.",
    "My Story: Derrick Watson’s professional journey, motivation, and the purpose of It’s Watts.",
]:
    bullet(doc, item)
callout(doc, "Out of scope for this release", "A full user-account system, progress tracking, paid courses, a community forum, and a CMS integration are not required for the initial future-updates release.")

h1(doc, "3. Functional requirements")
requirement(doc, "01", "What Is QA? page", "Create a beginner-friendly page explaining QA, its purpose, common responsibilities, and how it differs from simply finding bugs.", "Include plain-language sections for quality assurance, quality control, testing, the role of curiosity, and how QA supports users and teams.", "A new visitor can explain the purpose of QA and find a clear next action: start the beginner path, read the SDLC guide, or open the knowledge base.")
requirement(doc, "02", "QA in the SDLC guide", "Show where QA contributes from discovery and requirements through release and production feedback.", "Use a simple visual or step-by-step flow covering requirements, design, development, testing, release, and learning from production.", "The guide is readable on a narrow mobile screen without horizontal scrolling, and every phase includes at least one concrete QA contribution.")
requirement(doc, "03", "Knowledge Base", "Create a searchable reference area for foundational QA terms.", "Launch with 10–15 terms, including smoke testing, regression testing, exploratory testing, test case, defect, severity, priority, API, accessibility, CI/CD, and test plan.", "Visitors can search or filter terms and open a concise definition with a related learning-path or tool link.")
requirement(doc, "04", "Learning-level views", "Provide Beginner, Intermediate, and Advanced routes without requiring visitors to read every lesson.", "Each route needs a short description, 3–5 recommended topics, related tools/resources, and an explicit next step.", "The homepage exposes only three compact level cards. Each card opens a dedicated route or focused dialog; detailed lesson content does not all appear inline on the homepage.")
requirement(doc, "05", "Career Paths", "Help learners connect skill-building to realistic QA role progression.", "Include example routes such as Manual QA → Automation QA → Quality Engineer/SDET, while clearly stating that career paths are not one-size-fits-all.", "Each route lists key skills, portfolio practice ideas, and a recommended next resource or learning level.")
requirement(doc, "06", "Helpful Certifications", "Offer practical, non-gatekeeping certification guidance inside Career Paths.", "Present certifications as optional evidence of learning, not as replacements for practice. Include a foundation-testing option, accessibility-focused option, and role/tool-specific guidance. Link only to official certification information when published.", "The section clearly says that hands-on projects, testing judgment, communication, and product knowledge remain important. Certification details are kept concise and do not create a long home-page section.")
requirement(doc, "07", "QA orientation guide download", "Create a downloadable beginner guide that gives visitors a clear starting plan.", "Include: what QA is, essential terminology, a 30-day learning outline, first tools to explore, practice ideas, certification context, and links back to It’s Watts resources.", "The guide is downloadable from Career Paths and Resources, uses consistent It’s Watts branding, and contains no required paid recommendation.")
requirement(doc, "08", "My Story page", "Add a human, credible explanation of Derrick Watson’s QA journey and the mission of the site.", "Use the working page title My Story. Cover professional background, quality philosophy, why QA education matters, and how visitors can connect.", "The page remains concise, uses verified resume/portfolio information only, and can later be renamed Our Story if the brand becomes a broader team or community.")
requirement(doc, "09", "QA roadmap quiz", "Help visitors identify the most useful next learning level without requiring an account or a long assessment.", "Provide a short, mobile-friendly set of 4–6 questions about current experience, confidence, and learning goals. Return a recommended Beginner, Intermediate, or Advanced path with a clear next action.", "A visitor can complete the quiz in under two minutes, receive a reason for the recommendation, and open the matching learning path. The quiz must work without collecting personal data.")
requirement(doc, "10", "Glossary search", "Make key QA terminology easy to find and understand while building on the Knowledge Base.", "Add a prominent search or filter field, concise plain-language definitions, and links to related lessons, templates, or tools. Start with the planned foundational terms and allow the list to expand over time.", "A visitor can type a term or keyword and see relevant results immediately on desktop and mobile. Empty-result states suggest a related topic or a way to request a term.")
requirement(doc, "11", "Practice scenarios", "Turn passive reading into short, practical QA decision-making exercises.", "Provide compact scenario cards such as What would you test first?, What information belongs in this bug report?, or Should this be automated? Each includes a recommended answer and explanation after the visitor responds.", "At least three launch scenarios work with keyboard and touch input, show useful feedback without shaming incorrect selections, and fit within a focused destination rather than extending the homepage.")
requirement(doc, "12", "Expanded templates hub", "Give learners reusable starting points for real QA work.", "Extend the current downloads with test-case, bug-report, exploratory-testing, and release-checklist templates. Clearly describe when each template is useful and provide accessible download links.", "Visitors can preview each template’s purpose, download it successfully, and return to related learning content. Templates must use consistent It’s Watts branding and be useful without a paid tool or account.")
requirement(doc, "13", "Mobile dock regression safeguards", "Preserve reliable mobile navigation as focused learning destinations and future sections are added.", "Keep the dock visibly anchored to the bottom of the active mobile viewport. Reconnect scroll-based active-state tracking whenever a visitor opens or exits a focused view, such as Learn, so the current section remains highlighted.", "On a narrow device in portrait and landscape, the dock stays at the bottom while scrolling. Its active icon updates for Learn, Tools, Resources, Weekly Notes, Contact, and More; the behavior also works when the visitor starts on, enters, or exits a focused learning view.")

h1(doc, "4. Non-functional requirements")
for item in [
    "Mobile-first: no feature may require horizontal scrolling; home-page additions must remain compact and link to deeper destinations.",
    "Accessibility: keyboard navigation, visible focus, readable contrast in light and dark themes, labeled controls, and reduced-motion support must be retained.",
    "Performance: avoid large unoptimized media; load deeper content only when it is requested where practical.",
    "Theme support: every new panel, modal, guide, and interactive element must be tested in both system light and dark modes.",
    "Content governance: educational claims must be practical and clearly framed; certification details should be reviewed against official sources before publication.",
    "Responsive navigation: the mobile menu must provide direct access to Learn, Knowledge Base, Career Paths, Resources, and My Story without overwhelming the screen.",
    "Mobile dock regression checks: test viewport placement and active-section highlighting after any new focused route, conditional view, or navigation change.",
]:
    bullet(doc, item)

h1(doc, "5. Content and UX acceptance criteria")
for item in [
    "A first-time visitor can identify their learning level in under one minute.",
    "A mobile visitor can reach a deep-learning destination without scrolling through all future content on the homepage.",
    "Every Knowledge Base term is understandable without needing prior QA experience.",
    "Every career path includes a practical action, not only a list of credentials or job titles.",
    "The certification section avoids promising that any credential guarantees employment.",
    "The QA orientation guide can be found from at least two locations and downloaded successfully.",
    "All new pages retain the existing It’s Watts brand, light/dark theme behavior, and accessibility conventions.",
]:
    numbered(doc, item)

h1(doc, "6. Recommended delivery sequence")
for item in [
    "Phase 11A — Information architecture and foundation: add dedicated routes/navigation, What Is QA?, and QA in the SDLC.",
    "Phase 11B — Discovery and progression: add the Knowledge Base and Beginner / Intermediate / Advanced views.",
    "Phase 12 — Career enablement: add Career Paths, Helpful Certifications, and the QA orientation guide.",
    "Phase 13 — Brand depth: add My Story and refine internal links between learning, tools, resources, and weekly notes.",
    "Phase 14 — Practical learning tools: add the QA roadmap quiz, glossary search, practice scenarios, and expanded templates hub as the next focused release.",
]:
    numbered(doc, item)

h1(doc, "7. Next-release scope: practical learning tools")
callout(doc, "Recommendation", "Treat FR-09 through FR-12 as the next release. They directly support QA education, provide immediate practical value, and can live in focused destinations so the homepage remains short on mobile.")
label_bullet(doc, "Release goal", "Help a visitor identify where to start, understand unfamiliar terms, practice judgment, and leave with a reusable work template.")
label_bullet(doc, "Homepage rule", "Use one compact ‘Build your QA foundation’ teaser that links to the roadmap or learning hub. Do not add four full feature sections to the homepage.")
label_bullet(doc, "Later backlog", "Saved progress, email digests, public resource submissions, community or mentorship features, interactive accessibility scanning, and large tool-comparison collections should follow after content ownership, moderation, privacy, and maintenance decisions are made.")
label_bullet(doc, "Release checks", "Test each tool in system light and dark modes, on a narrow mobile screen, with keyboard navigation, and without an account or API key. Verify the mobile dock after opening and leaving every focused destination.")

h1(doc, "8. Dependencies and decisions needed later")
label_bullet(doc, "Content ownership", "Derrick provides or approves final educational copy, personal-story facts, and any certifications personally held or recommended.")
label_bullet(doc, "Certification accuracy", "Validate titles, requirements, costs, and availability against official providers at the time the content is published.")
label_bullet(doc, "Orientation guide format", "Start with a branded PDF or downloadable document; consider an email capture only after choosing a newsletter provider.")
label_bullet(doc, "CMS decision", "Keep the current code-based publishing workflow until weekly content volume makes a CMS worthwhile.")

h1(doc, "9. Success definition")
callout(doc, "Success", "It’s Watts becomes easier to navigate, not longer to scroll. Visitors can quickly understand QA, choose a learning level, find a career-oriented next step, and download a useful starting guide—on any screen size.")

doc.core_properties.title = "It’s Watts — Future Updates Requirements"
doc.core_properties.subject = "Requirements for QA learning-platform future updates"
doc.core_properties.author = "Derrick Watson"
doc.core_properties.keywords = "QA education, requirements, career paths, knowledge base, mobile-first"
doc.save(OUT)
print(OUT)
