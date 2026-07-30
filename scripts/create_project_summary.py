from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "Its_Watts_Project_Summary.docx"

INK = RGBColor(23, 35, 27)
GREEN = RGBColor(22, 132, 58)
MUTED = RGBColor(89, 102, 93)
LINE = RGBColor(213, 221, 215)
PALE_GREEN = "E8F3EB"


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill):
    props = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def add_bottom_border(paragraph, color="D5DDD7"):
    props = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "7")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    props.append(borders)


def add_page_field(paragraph):
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text.upper())
    set_font(run, 9, GREEN, True)
    run.font.all_caps = True
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 12.5, GREEN if level == 1 else INK, True)
    return p


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_font(run, 11, INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_font(run, 11, INK)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_font(run, 11, INK)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(8)
    shade_paragraph(p, PALE_GREEN)
    run = p.add_run(f"{label}: ")
    set_font(run, 10.5, GREEN, True)
    run = p.add_run(text)
    set_font(run, 10.5, INK)
    return p


def add_labeled_bullet(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(f"{label}. ")
    set_font(run, 11, INK, True)
    run = p.add_run(text)
    set_font(run, 11, INK)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.49)
section.footer_distance = Inches(0.49)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)

# Running header and footer.
header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
header.paragraph_format.space_after = Pt(4)
run = header.add_run("IT’S WATTS  |  PROJECT SUMMARY")
set_font(run, 8.5, MUTED, True)
add_bottom_border(header)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = footer.add_run("ITSWATTS.COM  |  PAGE ")
set_font(run, 8.5, MUTED)
add_page_field(footer)

# Cover / first-page furniture: editorial cover adapted to the site palette.
doc.add_paragraph().paragraph_format.space_after = Pt(42)
add_kicker(doc, "Website project summary")
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(9)
title_run = title.add_run("It’s Watts")
set_font(title_run, 31, INK, True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(16)
subtitle_run = subtitle.add_run("QA education, automation, and AI tools for practical quality work")
set_font(subtitle_run, 15, GREEN, True)

meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(34)
meta_run = meta.add_run("Project handoff and launch reference  •  July 2026  •  Built by Derrick Watson")
set_font(meta_run, 10.5, MUTED)
add_bottom_border(meta, "16843A")

add_callout(doc, "Purpose", "This document records the vision, design decisions, requirements, development process, representative prompts, and launch approach behind the It’s Watts website.")

add_heading(doc, "1. Overall goal and vision")
add_body(doc, "It’s Watts is designed as a practical QA education hub rather than a traditional portfolio. Its purpose is to help learners and working professionals make better quality decisions through approachable lessons, reusable resources, automation guidance, AI-assisted workflows, and small interactive tools.")
add_labeled_bullet(doc, "Brand promise", "Practical QA education for the AI era—clearer thinking, better testing habits, and software people can trust.")
add_labeled_bullet(doc, "Audience", "Early-career QA professionals, testers growing into automation, developers who want stronger quality habits, and teams looking for concise testing resources.")
add_labeled_bullet(doc, "Tone", "Direct, useful, calm, and credible. The site avoids exaggerated claims and frames AI as a thinking partner, not a replacement for professional judgment.")

add_heading(doc, "2. Site requirements and delivered capabilities")
add_body(doc, "The project started as a branded landing page and evolved into a learning-focused site with interactive education features.")
for item in [
    "A responsive React, Vite, and TypeScript website with reusable components.",
    "A branded glass-style navigation and modern QA-toolkit visual language.",
    "Learning paths for Automation Foundations, AI for QA, and Quality Essentials.",
    "Interactive QA tools: an automation decision guide, a context-aware test case starter, and a copyable QA prompt library.",
    "Weekly Notes cards with readable article dialogs and a visible content queue.",
    "Downloadable QA templates and a compact Resource Library with an expanded browse dialog.",
    "Ask Watt chat interface, connected through a secure server-side API route when API billing is enabled.",
    "System-preference light and dark themes, mobile navigation, accessibility support, social sharing metadata, robots.txt, and sitemap.xml.",
]:
    add_bullet(doc, item)

add_heading(doc, "3. Design direction")
add_body(doc, "The visual system intentionally resembles a focused engineering toolkit: dark charcoal foundations, clear hierarchy, technical grid details, green electric accents, thin borders, and light glass-like surfaces. The light theme preserves the same structure while using white panels, darker copy, and softer green accents.")
add_labeled_bullet(doc, "Core palette", "Dark #0E1116, panel #1B222B, accent green #3FB950, blue #58A6FF, and accessible light-theme equivalents.")
add_labeled_bullet(doc, "Typography", "Manrope for modern utility and readability, DM Mono for technical labels, and Playfair Display for selective editorial emphasis.")
add_labeled_bullet(doc, "Logo", "The displayed mark is it’s wattϟ. The lightning bolt acts as the stylized final “s,” while written brand references use It’s Watts.")
add_labeled_bullet(doc, "Interaction", "Cards lift subtly on hover, dialogs support Escape to close, focus states are visible, and reduced-motion preferences are respected.")

add_heading(doc, "4. Representative prompts and direction used")
add_body(doc, "The following summarizes the recurring creative and product direction used during development. These are not API prompts; they are the project brief that guided the site’s structure and content.")
for item in [
    "Build a polished personal technology brand around automation, AI, and quality assurance—not a conventional résumé-first portfolio.",
    "Make QA education the focus: include practical tools, automation guidance, AI-for-QA material, weekly notes, and reusable templates.",
    "Match the established portfolio palette: dark background, muted panels, thin technical borders, readable gray copy, and electric green accents.",
    "Use the brand It’s Watts in written content. Keep the lightning bolt after watt in the logo so it visually reads as the final “s.”",
    "Keep the page premium but practical: glass-like surfaces, purposeful animation, clear hierarchy, responsive mobile behavior, and no unnecessary visual clutter.",
    "Design Ask Watt as an education assistant that stays focused on QA, automation, accessibility, API testing, and responsible AI workflows.",
    "Make new content easy to publish from site data and downloadable files, then use Vercel for automatic production deployment.",
]:
    add_bullet(doc, item)

add_heading(doc, "5. Development process")
for step in [
    "Foundation. Established a Vite + React + TypeScript component architecture and a reusable CSS design system.",
    "Content focus. Shifted the primary experience from portfolio presentation to QA education, tools, learning paths, and weekly notes.",
    "Interactive learning. Added the automation guide, test-case starter, QA prompts, modal learning paths, and modal Weekly Notes.",
    "AI experience. Added Ask Watt’s client interface and a server-side Vercel API route that keeps the OpenAI API key private.",
    "Launch quality. Added metadata, social preview configuration, sitemap, robots.txt, keyboard support, reduced-motion support, responsive behavior, and light-mode support.",
    "Verification. Rebuilt the production bundle and ran lint checks after each significant change.",
]:
    add_numbered(doc, step)

add_heading(doc, "6. Why Vercel was used")
add_body(doc, "Vercel is a strong fit because It’s Watts is a Vite-based website that also needs a small secure backend endpoint for Ask Watt. It keeps the deployment workflow simple while avoiding the need to manage a separate server.")
for item in [
    "Automatic deployment from GitHub: pushing a change can create a new production deployment without manual file uploads.",
    "Global delivery network: the static site, templates, and images are served efficiently to visitors.",
    "Serverless API support: the /api/ask-watt route runs on the server, so OPENAI_API_KEY stays out of browser code and public repositories.",
    "Environment variables: the OpenAI key is stored in Vercel project settings rather than hard-coded into the website.",
    "Preview deployments: changes can be reviewed before going live, which supports a QA-minded workflow.",
    "Custom domain support: it connects the public site to itswatts.com while keeping deployment management separate from the domain registrar.",
]:
    add_bullet(doc, item)
add_callout(doc, "Important", "Vercel hosts and deploys the site. OpenAI API usage is billed separately. Ask Watt will respond only when the OpenAI project has an active key, available credits, and a permitted spending limit.")

add_heading(doc, "7. Content publishing workflow")
add_labeled_bullet(doc, "Weekly Notes", "Add a new entry to src/data/site.ts. The title, summary, intro, and article sections will appear automatically in the Weekly Notes area.")
add_labeled_bullet(doc, "Templates", "Place a Markdown file in public/downloads, then add its title and link in GrowthHub.tsx and ResourceLibrary.tsx.")
add_labeled_bullet(doc, "Deployment", "Commit and push the changes to GitHub. Vercel detects the update, builds the site, and publishes the new version.")
add_labeled_bullet(doc, "Future option", "A CMS such as Notion, Sanity, or another content platform can be added later if publishing without editing code becomes important.")

add_heading(doc, "8. Current status and next opportunities")
add_body(doc, "The site is launch-ready as a QA education hub. The next work is optional expansion rather than required foundation work.")
for item in [
    "Enable Ask Watt by confirming OpenAI billing and API credits, then validate the public chat endpoint after deployment.",
    "Publish the first real Weekly Note and one additional downloadable template.",
    "Choose a newsletter provider if email collection should move beyond the current mailto launch-list workflow.",
    "Choose analytics only if visitor measurement is needed; add it with a privacy-conscious configuration.",
    "Consider a CMS only when the manual content-file workflow becomes a constraint.",
]:
    add_bullet(doc, item)

add_heading(doc, "9. Final project principle")
add_callout(doc, "Guiding idea", "It’s Watts should feel like a useful place to learn—not a showcase that asks visitors to admire it. Every section should help someone make a clearer QA decision, build a stronger testing habit, or take a practical next step.")

doc.core_properties.title = "It’s Watts — Website Project Summary"
doc.core_properties.subject = "Vision, requirements, design, processes, prompts, and Vercel deployment rationale"
doc.core_properties.author = "Derrick Watson"
doc.core_properties.keywords = "QA education, automation, AI, Vercel, website project summary"
doc.save(OUT)
print(OUT)
